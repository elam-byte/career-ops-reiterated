#!/usr/bin/env python3
"""
Cover letter → .docx generator for Career Ops Reiterated.

Matches cv/coverletter.docx exactly:
  - Arial 11pt throughout
  - Justified text
  - Multi-line sender / recipient blocks
  - Bold section headings with exact spacing
  - List bullets with bold heading prefix + normal body
  - Spacing values from template (after: 0 / 5 / 6 / 9 / 10pt)

Body text format (Claude must output in this format):
  Normal paragraph  → plain text line
  Section heading   → **Heading text:**   (whole line wrapped in **)
  Bullet            → - **Bold heading:** rest of text

The salutation, subject, sign-off, and header blocks are passed as
separate flags so the script can format them independently.

Usage:
  python cv/generate_coverletter.py \\
      --company "Siemens" \\
      --role "Solutions Architect" \\
      --lang en \\
      --subject "Application as Solutions Architect" \\
      --salutation "Dear Ms. Müller," \\
      --hiring-manager "Ms. Müller" \\
      --company-address "Siemens AG\\nWerner-von-Siemens-Str. 1\\n80333 Munich" \\
      --text "Intro paragraph...\\n\\n**Section heading:**\\n..."

  # or pipe body via stdin
  echo "..." | python cv/generate_coverletter.py --company ... --role ... --lang ...

Requirements:
  pip install python-docx pyyaml
"""

import argparse
import re
import sys
from datetime import date
from pathlib import Path

import yaml

ROOT    = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "output"


# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

def load_profile() -> dict:
    profile_file = ROOT / "config" / "profile.yml"
    if profile_file.exists():
        with open(profile_file) as f:
            return yaml.safe_load(f) or {}
    return {}


def _make_slug(company: str, role: str) -> str:
    raw = f"{company}-{role[:40]}"
    raw = raw.lower().replace(" ", "-")
    return re.sub(r"[^a-z0-9-]", "", raw)


# ---------------------------------------------------------------------------
# Inline bold parser
# ---------------------------------------------------------------------------

def _parse_inline(text: str):
    """Split text on **bold** markers → list of (is_bold, chunk) tuples."""
    parts = []
    pattern = re.compile(r'\*\*(.+?)\*\*')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            parts.append((False, text[last:m.start()]))
        parts.append((True, m.group(1)))
        last = m.end()
    if last < len(text):
        parts.append((False, text[last:]))
    return parts or [(False, text)]


def _add_runs(para, text: str, size_pt: float = 11.0, bold_override: bool = False):
    """Add runs to a paragraph, honouring **bold** markers."""
    from docx.shared import Pt
    for is_bold, chunk in _parse_inline(text):
        run = para.add_run(chunk)
        run.bold = bold_override or is_bold
        run.font.name = "Arial"
        run.font.size = Pt(size_pt)


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------

def generate_docx(
    body_text: str,
    company: str,
    role: str,
    lang: str,
    profile: dict,
    subject_line: str = "",
    salutation: str = "",
    hiring_manager: str = "",
    company_address: str = "",
) -> Path:
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("ERROR: python-docx not installed. Run: pip install python-docx")
        sys.exit(1)

    candidate = profile.get("candidate", {})
    name    = candidate.get("full_name", "Elam Parithi Balasubramanian")
    email   = candidate.get("email",    "belamparithi22@gmail.com")
    phone   = candidate.get("phone",    "+49 176 28543935")
    address = candidate.get("location", "Leim 13, 88213 Oberzell \u2013 Ravensburg")

    doc = Document()

    # Page margins — match template
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # Default style
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    FONT = "Arial"
    SIZE = 11.0
    J    = WD_ALIGN_PARAGRAPH.JUSTIFY
    R    = WD_ALIGN_PARAGRAPH.RIGHT

    # ── Helper: add a simple paragraph ──────────────────────────────────────
    def add_para(text_content="", bold=False, after_pt=0, align=J):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(after_pt)
        if text_content:
            _add_runs(p, text_content, SIZE, bold_override=bold)
        return p

    # ── Helper: add a bullet paragraph ──────────────────────────────────────
    def add_bullet(text_content, after_pt=5):
        """List Bullet paragraph; text_content may contain **bold:** markers."""
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = J
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(after_pt)
        _add_runs(p, text_content, SIZE)
        return p

    # ── SENDER BLOCK ─────────────────────────────────────────────────────────
    add_para(name, bold=True, after_pt=0)
    add_para(address, after_pt=0)
    add_para(f"Telefon: {phone}", after_pt=0)
    add_para(f"E-Mail: {email}", after_pt=0)

    add_para(after_pt=0)  # blank line

    # ── RECIPIENT BLOCK ───────────────────────────────────────────────────────
    if hiring_manager:
        add_para(hiring_manager, bold=True, after_pt=0)
    for addr_line in company_address.split("\\n") if company_address else []:
        add_para(addr_line.strip(), after_pt=0)

    # ── DATE (right-aligned) ──────────────────────────────────────────────────
    if lang == "de":
        try:
            date_str = date.today().strftime("%-d. %B %Y")
        except ValueError:
            date_str = date.today().strftime("%d. %B %Y").lstrip("0")
        date_line = f"Ravensburg, {date_str}"
    else:
        date_line = f"Ravensburg, {date.today().strftime('%d %B %Y')}"

    add_para(date_line, align=R, after_pt=0)

    add_para(after_pt=0)  # blank line

    # ── SUBJECT LINE ──────────────────────────────────────────────────────────
    if subject_line:
        add_para(subject_line, bold=True, after_pt=10)

    # ── SALUTATION ────────────────────────────────────────────────────────────
    if salutation:
        add_para(salutation, after_pt=10)

    # ── BODY ──────────────────────────────────────────────────────────────────
    # Format rules:
    #   **Heading:**          → bold section heading, after=6pt
    #   - **Bold:** body text → list bullet with bold prefix, after=5pt
    #   - plain text          → list bullet, after=5pt
    #   (first line after heading) → bridge sentence, after=6pt
    #   (other lines)              → body paragraph, after=9pt

    lines = [l.rstrip() for l in body_text.strip().split("\n")]
    prev_type = None

    for line in lines:
        stripped = line.strip()

        if not stripped:
            continue

        # Section heading: **...**  on its own line (not a bullet)
        if (stripped.startswith("**") and stripped.endswith("**")
                and not stripped.startswith("- ") and not stripped.startswith("\u2022 ")):
            content = stripped[2:-2]
            add_para(content, bold=True, after_pt=6)
            prev_type = "heading"

        # Bullet with bold heading prefix: - **Bold:** body  OR  • **Bold:** body
        elif re.match(r'^[-\u2022]\s+\*\*', stripped):
            content = re.sub(r'^[-\u2022]\s+', '', stripped)
            add_bullet(content, after_pt=5)
            prev_type = "bullet"

        # Plain bullet: - text  OR  • text
        elif re.match(r'^[-\u2022]\s+', stripped):
            content = re.sub(r'^[-\u2022]\s+', '', stripped)
            p = doc.add_paragraph(style="List Bullet")
            p.alignment = J
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(5)
            r = p.add_run(content)
            r.font.name = FONT; r.font.size = Pt(SIZE)
            prev_type = "bullet"

        # Bridge sentence (first normal line after a heading)
        elif prev_type == "heading":
            add_para(stripped, after_pt=6)
            prev_type = "bridge"

        # Regular body paragraph
        else:
            add_para(stripped, after_pt=9)
            prev_type = "body"

    # ── CLOSING ───────────────────────────────────────────────────────────────
    add_para(after_pt=0)  # blank line before sign-off
    signoff = "Mit freundlichen Gr\u00fc\u00dfen" if lang == "de" else "Best regards,"
    add_para(signoff, after_pt=0)
    add_para(after_pt=0)
    add_para(name, after_pt=0)

    # ── SAVE ──────────────────────────────────────────────────────────────────
    slug      = _make_slug(company, role)
    build_dir = OUT_DIR / slug
    build_dir.mkdir(parents=True, exist_ok=True)
    out_path  = build_dir / f"coverletter-{slug}.docx"
    doc.save(str(out_path))
    return out_path


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Generate cover letter .docx")
    parser.add_argument("--company",         required=True,  help="Target company name")
    parser.add_argument("--role",            required=True,  help="Target role title")
    parser.add_argument("--lang",            choices=["en", "de"], default="en")
    parser.add_argument("--text",            help="Body text (or pipe via stdin)")
    parser.add_argument("--subject",         default="",     help="Subject line (bold)")
    parser.add_argument("--salutation",      default="",     help="Salutation line")
    parser.add_argument("--hiring-manager",  default="",     dest="hiring_manager",
                        help="Hiring manager name (bold in recipient block)")
    parser.add_argument("--company-address", default="",     dest="company_address",
                        help="Company address lines separated by \\n")
    args = parser.parse_args()

    if args.text:
        text = args.text
    elif not sys.stdin.isatty():
        text = sys.stdin.read()
    else:
        print("ERROR: Provide cover letter body via --text or stdin.")
        sys.exit(1)

    if not text.strip():
        print("ERROR: Cover letter body text is empty.")
        sys.exit(1)

    profile  = load_profile()
    out_path = generate_docx(
        body_text       = text,
        company         = args.company,
        role            = args.role,
        lang            = args.lang,
        profile         = profile,
        subject_line    = args.subject,
        salutation      = args.salutation,
        hiring_manager  = args.hiring_manager,
        company_address = args.company_address,
    )

    print(f"Cover letter saved: {out_path.resolve()}")
    print(f"Open with: libreoffice \"{out_path}\"")


if __name__ == "__main__":
    main()
